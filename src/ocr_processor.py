import os
import io
import base64
import logging
import requests
from PIL import Image
from key_manager import KeyManager


class AdvancedOCRWorker:
    def __init__(self, provider, api_key, formats, output_dir, custom_prompt="", example_text="", custom_model=None):
        self.provider = provider
        self.formats = formats
        self.output_dir = output_dir
        self.custom_prompt = custom_prompt
        self.example_text = example_text
        self.custom_model = custom_model

        # --- Gestione chiavi: Cassaforte > campo manuale ---
        km = KeyManager()
        km_keys = km.get_all_keys(provider)
        if km_keys:
            self.api_keys = km_keys
            logging.info(f"[OCR] Cassaforte: {len(km_keys)} chiave/i disponibili per {provider}.")
        elif api_key:
            self.api_keys = [api_key]
            logging.info(f"[OCR] Chiave manuale in uso per {provider}.")
        else:
            raise ValueError(f"Nessuna API Key disponibile per {provider}. Aprire la Cassaforte e inserire almeno una chiave.")

        self.current_key_idx = 0

    def _current_key(self):
        return self.api_keys[self.current_key_idx]

    def _rotate_key(self):
        """Passa alla prossima chiave disponibile. Ritorna False se ha esaurito il giro."""
        next_idx = (self.current_key_idx + 1) % len(self.api_keys)
        if next_idx == 0:
            return False  # Giro completo, tutte le chiavi sature
        self.current_key_idx = next_idx
        logging.warning(f"[OCR] Rotazione chiave → slot {next_idx + 1}/{len(self.api_keys)}")
        return True

    def process_file(self, f_path, review_callback=None, page_progress=None):
        """Elabora un singolo file (immagine o PDF) con retry automatico sulle chiavi."""
        logging.info(f"[OCR] Elaborazione: {os.path.basename(f_path)}")

        ext = os.path.splitext(f_path)[1].lower()
        if ext == ".pdf":
            self._process_pdf(f_path, review_callback, page_progress)
            return

        last_error = None
        for attempt in range(len(self.api_keys)):
            try:
                raw_text = self._transcribe_image(f_path, self._current_key())

                # Revisione interattiva opzionale
                final_text = raw_text
                if review_callback:
                    final_text = review_callback(f_path, raw_text)
                    if final_text is None:
                        logging.info(f"[OCR] Salvataggio saltato dall'utente.")
                        return

                self._save_results(f_path, final_text)
                return  # Successo

            except Exception as e:
                err_str = str(e).lower()
                last_error = e
                is_quota = any(k in err_str for k in ["429", "quota", "resource_exhausted", "rate"])
                if is_quota:
                    logging.warning(f"[OCR] Quota esaurita ({self._current_key()[:6]}...). Tentativo rotazione chiave...")
                    if not self._rotate_key():
                        break  # Tutte le chiavi sature
                else:
                    raise e  # Errore non recuperabile, rilancia

        raise Exception(f"[OCR] Tutte le chiavi esaurite. Ultimo errore: {last_error}")

    def _process_pdf(self, pdf_path, review_callback=None, page_progress=None):
        """Converte ogni pagina del PDF in immagine e le trascrive una per una."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise Exception("[OCR] PyMuPDF non installato. Esegui: pip install pymupdf")

        import tempfile
        doc = fitz.open(pdf_path)
        n_pages = len(doc)
        logging.info(f"[OCR] PDF: {n_pages} pagine trovate — {os.path.basename(pdf_path)}")

        all_texts = []
        with tempfile.TemporaryDirectory() as tmp_dir:
            for page_num in range(n_pages):
                if page_progress:
                    page_progress(page_num + 1, n_pages)

                page = doc[page_num]
                # Render a 200 DPI per leggibilità ottimale
                mat = fitz.Matrix(200 / 72, 200 / 72)
                pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
                img_path = os.path.join(tmp_dir, f"page_{page_num + 1:04d}.jpg")
                pix.save(img_path)
                logging.info(f"[OCR] PDF pagina {page_num + 1}/{n_pages} → {img_path}")

                # Trascrivi con retry chiavi
                last_error = None
                page_text = None
                for attempt in range(len(self.api_keys)):
                    try:
                        page_text = self._transcribe_image(img_path, self._current_key())
                        break
                    except Exception as e:
                        err_str = str(e).lower()
                        last_error = e
                        is_quota = any(k in err_str for k in ["429", "quota", "resource_exhausted", "rate"])
                        if is_quota:
                            if not self._rotate_key():
                                break
                        else:
                            raise e

                if page_text is None:
                    raise Exception(f"[OCR] Pagina {page_num + 1}: tutte le chiavi esaurite. {last_error}")

                # Revisione interattiva per pagina (immagine temp ancora disponibile)
                final_page_text = page_text
                if review_callback:
                    result = review_callback(img_path, page_text)
                    if result is None:
                        logging.info(f"[OCR] Pagina {page_num + 1} saltata dall'utente.")
                        continue
                    final_page_text = result

                all_texts.append(f"--- Pagina {page_num + 1} ---\n{final_page_text}")

        doc.close()

        if not all_texts:
            logging.info("[OCR] Nessuna pagina approvata. File non salvato.")
            return

        full_text = "\n\n".join(all_texts)
        self._save_results(pdf_path, full_text)

    def _prepare_image_b64(self, img_path):
        """Carica, ridimensiona e converte l'immagine in JPEG base64."""
        with Image.open(img_path) as img:
            max_size = 3072
            if max(img.size) > max_size:
                ratio = max_size / max(img.size)
                img = img.resize((int(img.size[0] * ratio), int(img.size[1] * ratio)), Image.Resampling.LANCZOS)
            if img.mode != "RGB":
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=90)
            return base64.b64encode(buf.getvalue()).decode("utf-8")

    def _build_prompt(self):
        """Costruisce il prompt da usare per la trascrizione."""
        if self.custom_prompt:
            prompt = self.custom_prompt
            if self.example_text:
                prompt += (
                    f"\n\nTRASCRIZIONE DI ESEMPIO (stessa calligrafia — segui questo stile):"
                    f"\n{self.example_text}\n\nINIZIA LA TRASCRIZIONE:"
                )
            return prompt
        prompt = (
            "Sei un esperto di paleografia e diplomatica.\n"
            "TRASCRIVI FEDELMENTE il testo contenuto in questa immagine.\n"
            "REGOLE ASSOLUTE:\n"
            "1. Mantieni la disposizione del testo originale (a capo, paragrafi).\n"
            "2. Non aggiungere commenti, introduzioni o conclusioni.\n"
            "3. Se il testo è incerto, usa [?].\n"
            "4. Sciogli le abbreviazioni ovvie tra parentesi angolate < >.\n"
            "5. Non normalizzare nomi propri o termini dialettali.\n"
        )
        if self.example_text:
            prompt += f"\nTRASCRIZIONE DI ESEMPIO (stessa calligrafia):\n{self.example_text}\n"
        prompt += "\nINIZIA LA TRASCRIZIONE:"
        return prompt

    def transcribe_top_preview(self, img_path, api_key, base_prompt):
        """Elabora solo il TOP 60% dell'immagine per la calibrazione assistita.
        Ritorna (tmp_image_path, text). Il chiamante deve eliminare tmp_image_path."""
        import tempfile

        with Image.open(img_path) as img:
            w, h = img.size
            top_pil = img.crop((0, 0, w, int(h * 0.60))).copy()

        # Ridimensiona e prepara b64
        max_size = 4096
        if max(top_pil.size) > max_size:
            ratio = max_size / max(top_pil.size)
            top_pil = top_pil.resize(
                (int(top_pil.size[0] * ratio), int(top_pil.size[1] * ratio)),
                Image.Resampling.LANCZOS
            )
        if top_pil.mode != "RGB":
            top_pil = top_pil.convert("RGB")

        # Salva crop come JPEG temporaneo (per mostrarlo nel dialog di revisione)
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".jpg", prefix="atk_calib_")
        os.close(tmp_fd)
        top_pil.save(tmp_path, format="JPEG", quality=95)

        buf = io.BytesIO()
        top_pil.save(buf, format="JPEG", quality=95)
        b64_top = base64.b64encode(buf.getvalue()).decode("utf-8")

        fmt_examples = (
            "\nESEMPIO FORMATO (RISPETTA ESATTAMENTE):"
            "\n  capofamiglia Ammogliato:  1 | 1 | 1 | Acciaj | Gaspero | 60 |  | 1 |  |  |  |  | Cattolico |  | Possidente | "
            "\n  moglie Maritata:          |  | 2 | \" | Rosa | 62 |  |  |  |  | 1 |  | \" |  |  | "
            "\n  figlio Celibe (età 3):    |  | 5 | \" | Giuseppe | 3 | 1 |  |  |  |  |  | \" |  |  | "
            "\n  figlia Nubile (età 4):    |  | 7 | \" | Rosa | 4 |  |  |  | 1 |  |  | \" |  |  | "
            "\n  figlia Nubile (età 1):    |  | 8 | \" | Tina | 1 |  |  |  | 1 |  |  | \" |  |  | "
            "\n  → ATTENZIONE: un bambino/a di 1 anno può essere sia maschio (Celibe, col.1) sia femmina (Nubile, col.4)."
            "\n  → LEGGI SEMPRE IL NOME ESATTAMENTE come scritto: non correggere 'Tina' in 'Pietro'."
            "\n  → Col.1 e Col.2 BLANK nelle righe di continuazione."
            "\n  → Il progressivo crescente va SEMPRE in Col.3."
            "\n⚠ STATO CIVILE — 6 caselle da sinistra: [1]Celibe [2]Ammogliato [3]Vedovo"
            " [4]Nubile [5]Maritata [6]Vedova. Segno in [1]=maschio Celibe; in [4]=femmina Nubile."
            " NON confondere casella [1] con [4]."
        )

        prompt_top = (
            base_prompt
            + "\n\nNOTA OPERATIVA — METÀ SUPERIORE (primo ~60% del foglio aperto):"
            " Trascrivi TUTTE le righe visibili, inclusa la riga di intestazione."
            " Le due pagine fisiche affiancate formano UN'UNICA RIGA: unisci le colonne con ' | '."
            "\n⚠ NON inventare righe non visibili nell'immagine."
            + fmt_examples
        )

        text = self._transcribe_gemini(api_key, b64_top, prompt_top, model=self.custom_model)
        logging.info(f"[OCR] Calibrazione: anteprima TOP estratta — {len(text.splitlines())} righe")
        return tmp_path, text

    def _transcribe_image(self, img_path, api_key):
        """Dispatcher: instrada al provider corretto.
        Per immagini a doppia pagina affiancata (aspect ratio > 1.6) con Gemini,
        esegue lo split TOP/BOTTOM per aumentare l'accuratezza OCR su tabelle lunghe."""
        prompt   = self._build_prompt()
        provider = (self.provider or "Gemini").strip()

        if "OpenAI" in provider or "GPT" in provider:
            return self._transcribe_openai(api_key, self._prepare_image_b64(img_path), prompt, model=self.custom_model)
        if "Anthropic" in provider or "Claude" in provider:
            return self._transcribe_claude(api_key, self._prepare_image_b64(img_path), prompt, model=self.custom_model)
        if "Mistral" in provider:
            return self._transcribe_openai_compat(api_key, img_path, prompt, "https://api.mistral.ai/v1", self.custom_model or "pixtral-large-latest")
        if "Groq" in provider:
            return self._transcribe_openai_compat(api_key, img_path, prompt, "https://api.groq.com/openai/v1", self.custom_model or "llama-3.2-90b-vision-preview")
        if "DeepSeek" in provider:
            # Solo testo — ignora immagine
            return self._transcribe_openai_compat(api_key, None, prompt, "https://api.deepseek.com", self.custom_model or "deepseek-chat")
        if "xAI" in provider:
            return self._transcribe_openai_compat(api_key, img_path, prompt, "https://api.x.ai/v1", self.custom_model or "grok-2-vision-1212")
        if "Ollama" in provider:
            host = api_key.strip() if api_key and api_key.strip().startswith("http") else "http://localhost:11434"
            return self._transcribe_openai_compat("ollama", img_path, prompt, host.rstrip("/") + "/v1", self.custom_model or "llava")
        if "Hugging" in provider or "HuggingFace" in provider:
            model = self.custom_model or "Qwen/Qwen2.5-VL-7B-Instruct"
            # TrOCR e modelli image-to-text puri: path diretto REST (testo grezzo)
            if any(m in model.lower() for m in ("trocr", "got-ocr", "olmocr", "pero-ocr")):
                return self._transcribe_hf_image_to_text(api_key, img_path, model)
            # Vision-language chat: endpoint OpenAI-compatibile
            return self._transcribe_openai_compat(api_key, img_path, prompt, "https://api-inference.huggingface.co/v1/", model)
        if "Transkribus" in provider:
            return self._transcribe_transkribus(api_key, img_path)

        # Default: Gemini con eventuale split doppia pagina
        with Image.open(img_path) as _probe:
            w, h = _probe.size
        is_double_page = (
            "DOPPIA PAGINA" in prompt
            or ((w / h) >= 1.6 if h > 0 else False)
        )
        if is_double_page:
            return self._transcribe_gemini_split(img_path, api_key, prompt)
        return self._transcribe_gemini(api_key, self._prepare_image_b64(img_path), prompt, model=self.custom_model)

    def _transcribe_gemini_split(self, img_path, api_key, prompt):
        """Split 2-vie (TOP/BOTTOM) con overlap 20% per immagini a doppia pagina.
        Risoluzione nativa: max_size=4096, quality=95.
        Merge: BOTTOM è master; TOP integra solo N°Casa/N°Fam se BOTTOM li ha vuoti."""

        with Image.open(img_path) as img:
            w, h = img.size
            # TOP: 0% → 60%,  BOTTOM: 40% → 100%  (overlap 20%)
            top_img = img.crop((0, 0,             w, int(h * 0.60)))
            bot_img = img.crop((0, int(h * 0.40), w, h))

            def _to_b64(pil_img):
                max_size = 4096
                if max(pil_img.size) > max_size:
                    ratio = max_size / max(pil_img.size)
                    pil_img = pil_img.resize(
                        (int(pil_img.size[0] * ratio), int(pil_img.size[1] * ratio)),
                        Image.Resampling.LANCZOS
                    )
                if pil_img.mode != "RGB":
                    pil_img = pil_img.convert("RGB")
                buf = io.BytesIO()
                pil_img.save(buf, format="JPEG", quality=95)
                return base64.b64encode(buf.getvalue()).decode("utf-8")

            b64_top = _to_b64(top_img)
            b64_bot = _to_b64(bot_img)

        fmt_examples = (
            "\nESEMPIO FORMATO (RISPETTA ESATTAMENTE):"
            "\n  capofamiglia Ammogliato:  1 | 1 | 1 | Acciaj | Gaspero | 60 |  | 1 |  |  |  |  | Cattolico |  | Possidente | "
            "\n  moglie Maritata:          |  | 2 | \" | Rosa | 62 |  |  |  |  | 1 |  | \" |  |  | "
            "\n  figlio Celibe (età 3):    |  | 5 | \" | Giuseppe | 3 | 1 |  |  |  |  |  | \" |  |  | "
            "\n  figlia Nubile (età 4):    |  | 7 | \" | Rosa | 4 |  |  |  | 1 |  |  | \" |  |  | "
            "\n  figlia Nubile (età 1):    |  | 8 | \" | Tina | 1 |  |  |  | 1 |  |  | \" |  |  | "
            "\n  → ATTENZIONE: un bambino/a di 1 anno può essere sia maschio (Celibe, col.1) sia femmina (Nubile, col.4)."
            "\n  → LEGGI SEMPRE IL NOME ESATTAMENTE come scritto: non correggere 'Tina' in 'Pietro'."
            "\n  → Col.1 e Col.2 BLANK nelle righe di continuazione."
            "\n  → Il progressivo crescente va SEMPRE in Col.3."
            "\n⚠ STATO CIVILE — 6 caselle da sinistra: [1]Celibe [2]Ammogliato [3]Vedovo"
            " [4]Nubile [5]Maritata [6]Vedova. Segno in [1]=maschio Celibe; in [4]=femmina Nubile."
            " NON confondere casella [1] con [4]."
        )

        prompt_top = (
            prompt
            + "\n\nNOTA OPERATIVA — METÀ SUPERIORE (primo ~60% del foglio aperto):"
            " Trascrivi TUTTE le righe visibili, inclusa la riga di intestazione."
            " Le due pagine fisiche affiancate formano UN'UNICA RIGA: unisci le colonne con ' | '."
            "\n⚠ NON inventare righe non visibili nell'immagine."
            + fmt_examples
        )

        logging.info("[OCR] Immagine doppia pagina — split TOP/BOTTOM (overlap 20%, max_size=4096, quality=95)")
        text_top = self._transcribe_gemini(api_key, b64_top, prompt_top, model=self.custom_model)

        # Contesto: ultime 3 righe del TOP per aiutare il BOTTOM a gestire i cambi nucleo
        _lines_ctx = [l for l in text_top.strip().splitlines() if l.strip()]
        _data_ctx  = _lines_ctx[1:] if (len(_lines_ctx) > 1 and any(
            k in _lines_ctx[0] for k in ("N°", "Cognome", "Nome", "Età", "Casa")
        )) else _lines_ctx
        _ctx_note = ""
        if _data_ctx[-3:]:
            _ctx_note = (
                "\n\nCONTESTO — ultime righe già trascritte dalla metà superiore:\n"
                + "\n".join(_data_ctx[-3:])
                + "\nSe la prima riga visibile qui inizia un NUOVO nucleo familiare"
                " (N°Casa e N°Famiglia diversi dall'ultima riga sopra),"
                " compilane obbligatoriamente le colonne 1 e 2 con i numeri letti nell'immagine."
            )

        prompt_bot = (
            prompt
            + "\n\nNOTA OPERATIVA — METÀ INFERIORE (ultimo ~60% del foglio aperto,"
            " con sovrapposizione rispetto alla metà superiore già trascritta):"
            " NON includere la riga di intestazione — inizia dalla prima riga dati."
            " Le due pagine fisiche affiancate formano UN'UNICA RIGA: unisci le colonne con ' | '."
            "\n⚠ NON inventare righe non visibili nell'immagine."
            + fmt_examples
            + _ctx_note
        )
        text_bot = self._transcribe_gemini(api_key, b64_bot, prompt_bot, model=self.custom_model)

        # Salva diagnostica raw per debugging
        try:
            import pathlib
            stem = pathlib.Path(img_path).stem
            diag_dir = self.output_dir if (self.output_dir and os.path.isdir(self.output_dir)) else os.getcwd()
            for label, raw in (("TOP", text_top), ("BOTTOM", text_bot)):
                diag_path = os.path.join(diag_dir, f"DIAG_{stem}_{label}.txt")
                with open(diag_path, "w", encoding="utf-8") as _df:
                    _df.write(raw)
            logging.debug(f"[OCR] Diagnostica split salvata in: {diag_dir}")
        except Exception as _de:
            logging.warning(f"[OCR] Impossibile salvare diagnostica split: {_de}")

        # --- Merge: BOTTOM master, TOP integra solo N°Casa/N°Fam se BOTTOM li ha vuoti ---
        def _get_prog(row: str):
            parts = [p.strip() for p in row.split("|")]
            if len(parts) >= 3:
                try:
                    return int(parts[2])
                except ValueError:
                    pass
            return None

        def _has_header(line: str) -> bool:
            return any(k in line for k in ("N°", "Religione", "Cognome", "Nome", "Età", "Casa", "Famiglia"))

        lines_top = [l for l in text_top.strip().splitlines() if l.strip()]
        lines_bot = [l for l in text_bot.strip().splitlines() if l.strip()]

        # Rimuovi intestazione dal bottom se il modello l'ha inclusa
        if lines_bot and _has_header(lines_bot[0]):
            lines_bot = lines_bot[1:]

        if not lines_top:
            return "\n".join(lines_bot)

        header_lines = [lines_top[0]] if lines_top else []
        data_top = lines_top[1:] if (lines_top and _has_header(lines_top[0])) else lines_top

        # Mappa prog → riga per il TOP
        top_map: dict[int, str] = {}
        for row in data_top:
            p = _get_prog(row)
            if p is not None:
                top_map[p] = row

        valid_bot_progs = [p for p in (_get_prog(l) for l in lines_bot) if p is not None]
        first_bot_prog  = min(valid_bot_progs) if valid_bot_progs else None
        max_top_prog    = max(top_map.keys())  if top_map        else None

        def _merge_cols(top_row: str, bot_row: str) -> str:
            """BOTTOM master: integra N°Casa(col 0) e N°Fam(col 1) dal TOP solo se BOTTOM li ha vuoti."""
            tc = [c.strip() for c in top_row.split("|")]
            bc = [c.strip() for c in bot_row.split("|")]
            for idx in (0, 1):
                if idx < len(tc) and idx < len(bc) and not bc[idx] and tc[idx]:
                    bc[idx] = tc[idx]
            return " | ".join(bc)

        if first_bot_prog is not None and max_top_prog is not None:
            rows_top_only   = [r for r in data_top   if (_get_prog(r) or 0) < first_bot_prog]
            overlap_rows    = []
            for row in lines_bot:
                p = _get_prog(row)
                if p is not None and p <= max_top_prog:
                    overlap_rows.append(_merge_cols(top_map[p], row) if p in top_map else row)
            rows_bot_only   = [r for r in lines_bot  if (_get_prog(r) or 0) > max_top_prog]
            merged_data = rows_top_only + overlap_rows + rows_bot_only
        else:
            top_only = [r for r in data_top if (_get_prog(r) or 0) < (first_bot_prog or float("inf"))]
            merged_data = top_only + lines_bot

        return "\n".join(header_lines + merged_data)

    def _call_gemini_rest(self, api_key, model, b64_img, prompt):
        """Chiama l'API REST Gemini con un'immagine e ritorna il testo."""
        url = f"https://generativelanguage.googleapis.com/v1/models/{model}:generateContent?key={api_key}"
        payload = {
            "contents": [{"parts": [
                {"text": prompt},
                {"inline_data": {"mime_type": "image/jpeg", "data": b64_img}}
            ]}],
            "generationConfig": {"temperature": 0.1, "maxOutputTokens": 32768}
        }
        resp = requests.post(url, headers={"Content-Type": "application/json"}, json=payload, timeout=180)
        resp.raise_for_status()
        res_json = resp.json()
        if "candidates" in res_json:
            candidate = res_json["candidates"][0]
            finish_reason = candidate.get("finishReason", "UNKNOWN")
            content = candidate.get("content", {})
            parts = content.get("parts", [])
            if parts and "text" in parts[0]:
                return parts[0]["text"]
            logging.warning(f"[OCR] Gemini risposta senza testo. finishReason={finish_reason}, content={str(content)[:200]}")
            raise Exception(f"Risposta Gemini senza testo (finishReason={finish_reason})")
        raise Exception(f"Risposta Gemini vuota: {res_json}")

    def _transcribe_gemini(self, api_key, b64_img, prompt, model=None):
        """Trascrizione via Gemini REST. Per OCR usa pro → flash come fallback."""
        from ai_utils import get_best_gemini_model

        # Ordine di preferenza: prima pro (più accurato su tabelle), poi flash
        if model:
            model_candidates = [model]
        else:
            model_candidates = [
                get_best_gemini_model(api_key, preferred="pro"),
                get_best_gemini_model(api_key, preferred="flash"),
            ]
        # Rimuovi duplicati mantenendo l'ordine
        seen = set()
        model_candidates = [m for m in model_candidates if m not in seen and not seen.add(m)]

        last_error = None
        for model_name in model_candidates:
            try:
                logging.info(f"[OCR] Gemini — modello: {model_name}")
                return self._call_gemini_rest(api_key, model_name, b64_img, prompt)
            except Exception as e:
                err_str = str(e).lower()
                last_error = e
                if any(k in err_str for k in ["404", "not found", "does_not_exist", "invalid"]):
                    logging.warning(f"[OCR] Modello {model_name} non disponibile, provo successivo.")
                    continue
                raise e
        raise Exception(f"[OCR] Nessun modello Gemini disponibile. Ultimo errore: {last_error}")

    def _transcribe_openai(self, api_key, b64_img, prompt, model=None):
        """Trascrizione via OpenAI Vision (gpt-4o)."""
        if not model: model = 'gpt-4o'
        logging.info(f"[OCR] OpenAI — modello: {model}")
        url = "https://api.openai.com/v1/chat/completions"
        payload = {
            "model": model,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url",
                     "image_url": {"url": f"data:image/jpeg;base64,{b64_img}", "detail": "high"}}
                ]
            }],
            "temperature": 0.1,
            "max_tokens": 4096
        }
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        resp = requests.post(url, headers=headers, json=payload, timeout=120)
        resp.raise_for_status()
        res_json = resp.json()
        try:
            return res_json["choices"][0]["message"]["content"]
        except (KeyError, IndexError):
            raise Exception(f"Risposta OpenAI non valida: {res_json}")

    def _transcribe_claude(self, api_key, b64_img, prompt, model=None):
        """Trascrizione via Anthropic Messages API."""
        if not model: model = "claude-3-5-sonnet-latest"
        logging.info(f"[OCR] Anthropic — modello: {model}")
        url = "https://api.anthropic.com/v1/messages"
        payload = {
            "model": model,
            "max_tokens": 4096,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image",
                     "source": {"type": "base64", "media_type": "image/jpeg", "data": b64_img}},
                    {"type": "text", "text": prompt}
                ]
            }]
        }
        headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        }
        resp = requests.post(url, headers=headers, json=payload, timeout=120)
        resp.raise_for_status()
        res_json = resp.json()
        try:
            return res_json["content"][0]["text"]
        except (KeyError, IndexError):
            raise Exception(f"Risposta Anthropic non valida: {res_json}")

    def _transcribe_openai_compat(self, api_key, img_path, prompt, base_url, model):
        """Trascrizione via endpoint OpenAI-compatibile (Mistral, Groq, xAI, Ollama, HuggingFace chat)."""
        from openai import OpenAI
        logging.info(f"[OCR] OpenAI-compat — base_url: {base_url} modello: {model}")
        client = OpenAI(api_key=api_key, base_url=base_url)
        content = [{"type": "text", "text": prompt}]
        if img_path and os.path.exists(img_path):
            b64 = self._prepare_image_b64(img_path)
            content.insert(0, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
        response = client.chat.completions.create(
            model=model,
            max_tokens=8192,
            messages=[{"role": "user", "content": content}]
        )
        return response.choices[0].message.content

    def _transcribe_hf_image_to_text(self, api_key, img_path, model):
        """Trascrizione via HuggingFace Inference API image-to-text (TrOCR, GOT-OCR2, ecc.).
        Restituisce testo grezzo — adatto come input per la fase GEDCOM."""
        logging.info(f"[OCR] HuggingFace image-to-text — modello: {model}")
        url = f"https://api-inference.huggingface.co/models/{model}"
        headers = {"Authorization": f"Bearer {api_key}"}
        b64 = self._prepare_image_b64(img_path)
        img_bytes = base64.b64decode(b64)
        resp = requests.post(url, headers={**headers, "Content-Type": "image/jpeg"},
                             data=img_bytes, timeout=120)
        resp.raise_for_status()
        result = resp.json()
        if isinstance(result, list) and result:
            return result[0].get("generated_text", str(result[0]))
        if isinstance(result, dict):
            return result.get("generated_text", str(result))
        return str(result)

    def _transcribe_transkribus(self, api_key, img_path):
        """Trascrizione via Transkribus HTR REST API.
        
        Credenziali (campo API Key):
          - Formato 'email:password': autenticazione classica con JSESSIONID
          - Formato token bearer: qualsiasi stringa senza ':' (es. chiave API generata da account)
        
        Modello HTR (campo Override Modello):
          - Lascia vuoto  → 48122 (Italian Handwriting M1, modello ufficiale Archivi di Stato)
          - Inserisci ID numerico → usa quel modello specifico
        
        Il documento temporaneo caricato su Transkribus viene rimosso al termine.
        """
        import time
        import xml.etree.ElementTree as ET

        TRANSKRIBUS_BASE = "https://transkribus.eu/TrpServer/rest"
        model_id = int(self.custom_model) if (self.custom_model or "").strip().isdigit() else 48122
        logging.info(f"[OCR] Transkribus — modello HTR: {model_id}, file: {os.path.basename(img_path)}")

        # ── 1. Autenticazione ─────────────────────────────────────────────────
        if ":" in api_key:
            # email:password → session cookie
            email, password = api_key.split(":", 1)
            r_auth = requests.post(
                f"{TRANSKRIBUS_BASE}/auth/login",
                data={"user": email.strip(), "pw": password.strip()},
                headers={"Content-Type": "application/x-www-form-urlencoded"},
                timeout=30,
            )
            r_auth.raise_for_status()
            session_id = (r_auth.json() or {}).get("sessionId")
            if not session_id:
                raise Exception(f"[Transkribus] Login fallito. Risposta: {r_auth.text[:300]}")
            auth_headers = {"Cookie": f"JSESSIONID={session_id}"}
            logging.info("[OCR] Transkribus auth via email:password — OK")
        else:
            # Bearer token (API token generato dall'account)
            auth_headers = {"Authorization": f"Bearer {api_key.strip()}"}
            logging.info("[OCR] Transkribus auth via Bearer token")

        # ── 2. Prima collection disponibile ───────────────────────────────────
        r_coll = requests.get(
            f"{TRANSKRIBUS_BASE}/collections/list",
            headers=auth_headers, timeout=30,
        )
        r_coll.raise_for_status()
        collections = r_coll.json()
        if not collections:
            raise Exception(
                "[Transkribus] Nessuna collection trovata nell'account.\n"
                "Crea almeno una collection su transkribus.eu prima di usare questo provider."
            )
        coll_id = collections[0]["colId"]
        logging.info(f"[OCR] Transkribus — collection: {coll_id} ({collections[0].get('colName', '')})")

        # ── 3. Upload immagine → crea documento temporaneo ───────────────────
        mime = "image/png" if img_path.lower().endswith(".png") else "image/jpeg"
        basename = os.path.basename(img_path)
        with open(img_path, "rb") as fh:
            img_bytes = fh.read()

        r_upload = requests.post(
            f"{TRANSKRIBUS_BASE}/collections/{coll_id}/createDocFromImages",
            headers=auth_headers,
            files={"file": (basename, img_bytes, mime)},
            data={"title": f"ATKPro-tmp-{basename}"},
            timeout=90,
        )
        r_upload.raise_for_status()
        upload_data = r_upload.json()
        doc_id = upload_data if isinstance(upload_data, int) else upload_data.get("docId")
        if not doc_id:
            raise Exception(f"[Transkribus] Upload fallito: {str(upload_data)[:200]}")
        logging.info(f"[OCR] Transkribus — documento temporaneo creato: docId={doc_id}")

        # ── 4. Avvia job HTR ──────────────────────────────────────────────────
        htr_payload = {
            "collId": coll_id,
            "docList": {
                "docs": [{"docId": doc_id, "pageList": {"pages": [{"pageNr": 1}]}}]
            },
            "htrId": model_id,
        }
        r_htr = requests.post(
            f"{TRANSKRIBUS_BASE}/recognition/{coll_id}/htrRnn",
            json=htr_payload,
            headers={**auth_headers, "Content-Type": "application/json"},
            timeout=60,
        )
        r_htr.raise_for_status()
        job_data = r_htr.json()
        job_id = job_data if isinstance(job_data, int) else job_data.get("jobId") or job_data.get("jobIds", [None])[0]
        if not job_id:
            raise Exception(f"[Transkribus] Avvio HTR fallito: {str(job_data)[:200]}")
        logging.info(f"[OCR] Transkribus — job HTR avviato: jobId={job_id}")

        # ── 5. Polling fino al completamento (max 180s) ───────────────────────
        max_wait, poll_interval, elapsed = 180, 6, 0
        final_state = None
        while elapsed < max_wait:
            time.sleep(poll_interval)
            elapsed += poll_interval
            try:
                r_status = requests.get(
                    f"{TRANSKRIBUS_BASE}/jobs/{job_id}",
                    headers=auth_headers, timeout=30,
                )
                if r_status.status_code != 200:
                    continue
                state = str(r_status.json().get("state", "")).upper()
                logging.debug(f"[OCR] Transkribus job {job_id} — stato: {state} ({elapsed}s)")
                if state == "FINISHED":
                    final_state = state
                    break
                if state in ("FAILED", "CANCELED"):
                    desc = r_status.json().get("description", "")
                    raise Exception(f"[Transkribus] Job {job_id} {state}: {desc}")
            except requests.RequestException:
                pass  # Retry al prossimo ciclo
        if final_state != "FINISHED":
            raise Exception(f"[Transkribus] Timeout ({max_wait}s) attendendo il job {job_id}.")

        # ── 6. Recupera PAGE XML della pagina trascritta ──────────────────────
        r_page = requests.get(
            f"{TRANSKRIBUS_BASE}/collections/{coll_id}/{doc_id}/1/curr",
            headers=auth_headers, timeout=30,
        )
        r_page.raise_for_status()
        page_xml = r_page.text

        # ── 7. Estrazione testo da PAGE XML ───────────────────────────────────
        text_lines = []
        try:
            root = ET.fromstring(page_xml)
            # Supporta entrambe le versioni del namespace PAGE XML
            for ns_uri in (
                "http://schema.primaresearch.org/PAGE/gts/pagecontent/2013-07-15",
                "http://schema.primaresearch.org/PAGE/gts/pagecontent/2019-07-15",
            ):
                tag = f"{{{ns_uri}}}Unicode"
                nodes = root.iter(tag)
                lines = [n.text for n in nodes if n.text and n.text.strip()]
                if lines:
                    text_lines = lines
                    break
        except ET.ParseError:
            pass
        if not text_lines:
            # Fallback grezzo — estrae testo tra tag Unicode
            import re as _re
            text_lines = _re.findall(r"<Unicode>(.*?)</Unicode>", page_xml, _re.DOTALL)
        result_text = "\n".join(text_lines)
        logging.info(f"[OCR] Transkribus — {len(text_lines)} righe estratte")

        # ── 8. Pulizia documento temporaneo (best-effort) ─────────────────────
        try:
            requests.delete(
                f"{TRANSKRIBUS_BASE}/collections/{coll_id}/{doc_id}",
                headers=auth_headers, timeout=20,
            )
            logging.info(f"[OCR] Transkribus — documento temporaneo {doc_id} rimosso")
        except Exception as _cleanup_err:
            logging.warning(f"[OCR] Transkribus — pulizia doc {doc_id} fallita (ignorata): {_cleanup_err}")

        return result_text

    def _save_results(self, orig_path, text):
        base_name = os.path.splitext(os.path.basename(orig_path))[0]

        if "txt" in self.formats:
            out_p = os.path.join(self.output_dir, f"{base_name}_trascrizione.txt")
            with open(out_p, "w", encoding="utf-8") as f:
                f.write(text)
            logging.info(f"[OCR] TXT salvato: {out_p}")

        if "docx" in self.formats:
            try:
                import docx
                doc = docx.Document()
                doc.add_heading(f"Trascrizione: {base_name}", 0)
                doc.add_paragraph(text)
                out_p = os.path.join(self.output_dir, f"{base_name}_trascrizione.docx")
                doc.save(out_p)
                logging.info(f"[OCR] DOCX salvato: {out_p}")
            except Exception as e:
                logging.error(f"[OCR] Errore salvataggio DOCX: {e}")

        if "xml" in self.formats:
            out_p = os.path.join(self.output_dir, f"{base_name}_trascrizione.xml")
            from datetime import date as _date
            tei_body = self._text_to_tei_body(text)
            xml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<TEI xmlns="http://www.tei-c.org/ns/1.0">\n'
                '  <teiHeader>\n'
                '    <fileDesc>\n'
                '      <titleStmt>\n'
                f'        <title>Trascrizione di {base_name}</title>\n'
                '      </titleStmt>\n'
                '      <publicationStmt>\n'
                f'        <p>Trascrizione automatica ATK-Pro — {_date.today().isoformat()}</p>\n'
                '      </publicationStmt>\n'
                '      <sourceDesc>\n'
                f'        <p>File originale: {os.path.basename(orig_path)}</p>\n'
                '      </sourceDesc>\n'
                '    </fileDesc>\n'
                '  </teiHeader>\n'
                '  <text>\n'
                '    <body>\n'
                '      <div type="transcription">\n'
                f'{tei_body}'
                '      </div>\n'
                '    </body>\n'
                '  </text>\n'
                '</TEI>'
            )
            with open(out_p, "w", encoding="utf-8") as f:
                f.write(xml)
            logging.info(f"[OCR] XML TEI-P5 salvato: {out_p}")

    def _text_to_tei_body(self, text):
        """Converte il testo trascritto in markup TEI-P5 valido."""
        import re
        out_lines = []
        for line in text.split('\n'):
            # --- Pagina N --- → <pb n="N"/>
            m = re.match(r'^---\s*Pagina\s+(\d+)\s*---$', line.strip())
            if m:
                out_lines.append(f'        <pb n="{m.group(1)}"/>')
                continue

            # Estrai le espansioni < > PRIMA di fare l'escaping
            parts = []
            pos = 0
            for am in re.finditer(r'<([^>]+)>', line):
                before = line[pos:am.start()].replace('&', '&amp;')
                parts.append(before)
                parts.append(f'<expan>{am.group(1)}</expan>')
                pos = am.end()
            remaining = line[pos:].replace('&', '&amp;')
            parts.append(remaining)
            line_content = ''.join(parts)

            # [?] → <unclear reason="illegible"/>
            line_content = re.sub(r'\[\?\]', '<unclear reason="illegible"/>', line_content)

            out_lines.append(f'        <ab>{line_content}</ab>')

        return '\n'.join(out_lines) + '\n'
